# -*- coding: utf-8 -*-
"""Genere le resume executif (5 pages max) au format Word — charte graphique AfriMarket."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG = r"c:\Users\LENOVO\Desktop\PROJET PYTHON\figures\brand"
LOGO = r"c:\Users\LENOVO\Desktop\PROJET PYTHON\assets\logo.png"
OUT = r"c:\Users\LENOVO\Desktop\PROJET PYTHON\reports\Resume_Executif_AfriMarket.docx"

# --- Charte graphique AfriMarket (extraite du logo) ---
RED = RGBColor(0xC4, 0x2D, 0x1C)
RED_HEX = "C42D1C"
GRAY = RGBColor(0x3F, 0x3F, 0x3F)
GRAY_LIGHT = RGBColor(0x8C, 0x8C, 0x8C)
GRAY_PALE_HEX = "F2F2F2"

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)


def set_cell_shading(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = GRAY
    return p


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def add_image_row(paths, width_each=2.3):
    table = doc.add_table(rows=1, cols=len(paths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, path in enumerate(paths):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_each))


# =====================================================================
# EN-TETE DE PAGE (logo + libelle) — pages 2 et suivantes uniquement
# (la page 1 a deja son propre bandeau logo/titre dans le corps du document)
# =====================================================================
doc.sections[0].different_first_page_header_footer = True
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run()
hr.add_picture(LOGO, height=Cm(0.9))
hp2 = header.add_paragraph()
hp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp2.add_run("RÉSUMÉ EXÉCUTIF — ANALYSE STRATÉGIQUE")
r.font.size = Pt(8)
r.font.color.rgb = GRAY_LIGHT
r.bold = True
# Filet rouge sous l'en-tete
p_border = OxmlElement("w:pPr")
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:space"), "4")
bottom.set(qn("w:color"), RED_HEX)
pbdr.append(bottom)
hp2._p.get_or_add_pPr().append(pbdr)

# =====================================================================
# PAGE 1 — TITRE + CONTEXTE + METHODOLOGIE + KPIS
# =====================================================================
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.add_run().add_picture(LOGO, height=Cm(2.4))

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(8)
r = title.add_run("RÉSUMÉ EXÉCUTIF — ANALYSE STRATÉGIQUE")
r.bold = True
r.font.size = Pt(19)
r.font.color.rgb = GRAY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Performance commerciale · Juillet – Décembre 2025")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RED

doc.add_paragraph()

h1("Contexte & objectif")
body(
    "AfriMarket est une entreprise e-commerce panafricaine (8 villes, 4 catégories : Électronique, "
    "Mode, Beauté, Maison) confrontée à des variations de chiffre d'affaires, un taux de retour "
    "préoccupant et des écarts de performance selon les villes. Cette analyse porte sur 6 mois "
    "d'activité (9 400 commandes valides après nettoyage) et vise à éclairer les décisions de la "
    "direction sur les catégories, la géographie, le marketing et les clients."
)

h1("Méthodologie")
body(
    "Le dataset brut (10 100 lignes) contenait des doublons (100), des prix aberrants (622 valeurs "
    "négatives imputées par la médiane de catégorie), des remises négatives (600, corrigées), des "
    "quantités nulles (600, supprimées), des villes mal orthographiées (Kinshassa) et des catégories "
    "incohérentes (electronique / Électronique). Après nettoyage, le dataset final compte 9 400 "
    "commandes exploitables, enrichies d'indicateurs métier (CA, marge brute et profit net estimés "
    "par une hypothèse de taux de marge par catégorie, taux de retour, valeur vie client)."
)

h1("Indicateurs clés de performance (KPIs)")

kpi_table = doc.add_table(rows=2, cols=6)
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["CA total", "Profit net est.", "Panier moyen", "Taux annulation", "Taux retour", "Clients actifs"]
values = ["2 507 325 $", "409 924 $", "272 $", "1.95 %", "8.14 %", "1 747"]
for i, htxt in enumerate(headers):
    cell = kpi_table.rows[0].cells[i]
    cell.text = htxt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, RED_HEX)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for i, v in enumerate(values):
    cell = kpi_table.rows[1].cells[i]
    cell.text = v
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = GRAY
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, GRAY_PALE_HEX)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "Lecture : le CA mensuel oscille sans tendance nette (410k$–460k$), confirmant l'instabilité "
    "constatée par la direction — expliquée en détail par catégorie et par ville ci-dessous."
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GRAY_LIGHT

doc.add_page_break()

# =====================================================================
# PAGE 2 — CATEGORIE + GEOGRAPHIE
# =====================================================================
h1("Analyse par catégorie — Quelle catégorie prioriser ou optimiser ?")
add_image_row([f"{FIG}/01_ca_categorie.png", f"{FIG}/02_retour_categorie.png"], width_each=2.9)
body(
    "Électronique génère 74,6 % du CA mais cumule la marge la plus faible (hypothèse 15 %) et le "
    "taux de retour le plus élevé (13,8 %) : c'est la catégorie à optimiser en priorité "
    "(qualité produit, sélection fournisseurs). Mode et Beauté offrent des marges de 45 à 50 % avec "
    "un taux de retour sous 7,5 % mais ne pèsent que 10,4 % du CA cumulé : ce sont les catégories à "
    "prioriser pour une croissance rentable."
)

h1("Analyse géographique — Où investir davantage ?")
add_image_row([f"{FIG}/04_ca_ville.png", f"{FIG}/05_annulation_ville.png"], width_each=2.9)
body(
    "Kinshasa, Abidjan et Dakar sont les marchés les plus solides. Douala affiche la plus forte "
    "croissance (+76 % sur la période) mais aussi le taux d'annulation le plus élevé (12,9 %, contre "
    "0-1 % ailleurs) — signe d'un problème opérationnel local à résoudre avant d'y investir "
    "davantage. Libreville recule (-46 %) et mérite un diagnostic dédié."
)

doc.add_page_break()

# =====================================================================
# PAGE 3 — MARKETING + CLIENTS
# =====================================================================
h1("Analyse marketing — Quel canal mérite plus de budget ?")
add_image_row([f"{FIG}/06_roi_marketing.png"], width_each=4.5)
body(
    "Email affiche de loin le meilleur ROI (231x) pour un budget marginal (2,3k$) : canal "
    "sous-investi à scaler en priorité. Instagram Ads génère le plus de CA et la meilleure "
    "rétention (54 %) mais avec un ROI plus modéré (24,7x) du fait d'un budget déjà conséquent. "
    "Influenceur cumule le ROI le plus faible (21,6x) et la rétention la plus faible (42 %) : "
    "candidat à la réduction budgétaire."
)

h1("Analyse clients — Comment améliorer la rétention ?")
add_image_row([f"{FIG}/07_pareto_clients.png", f"{FIG}/08_segmentation_clients.png"], width_each=2.9)
body(
    "74 % des clients sont déjà récurrents (≥2 commandes) et 32 % des clients génèrent 80 % du CA. "
    "Le principal levier de croissance rentable est la conversion des 26 % de clients à commande "
    "unique en acheteurs récurrents, via le canal Email (meilleur ROI identifié)."
)

doc.add_page_break()

# =====================================================================
# PAGE 4 — SYNTHESE + RECOMMANDATIONS
# =====================================================================
h1("Synthèse des insights clés")
bullet("Électronique = 74,6 % du CA mais marge la plus faible (15 %) et taux de retour le plus élevé (13,8 %).")
bullet("Douala : plus forte croissance (+76 %) ET plus fort taux d'annulation (12,9 %) — signal d'alerte opérationnel.")
bullet("Email : ROI le plus élevé (231x) mais budget marginal — sous-investi.")
bullet("32 % des clients génèrent 80 % du CA ; 26 % des clients n'ont commandé qu'une fois.")
bullet("Profit net global estimé (~410k$) nettement inférieur au potentiel du CA brut (~2,51M$), du fait de la marge structurellement faible de l'électronique.")

h1("5 recommandations stratégiques")

recos = [
    ("1", "Réduire le taux de retour sur l'Électronique (13,8 %)",
     "Auditer la qualité produit et les fiches descriptives des références les plus retournées ; "
     "renforcer le contrôle qualité fournisseur. Impact démultiplié sur le profit net vu le volume."),
    ("2", "Corriger le problème opérationnel à Douala avant d'y investir davantage",
     "Diagnostiquer la cause du taux d'annulation de 12,9 % (paiement, délais de livraison) pour ne "
     "pas acquérir des clients qui annulent immédiatement malgré la forte croissance du marché."),
    ("3", "Réallouer le budget marketing : renforcer Email, réduire Influenceur",
     "Multiplier les campagnes Email (ROI 231x, budget quasi nul) ; réduire ou renégocier les "
     "partenariats Influenceur (ROI 21,6x, rétention la plus faible : 42 %)."),
    ("4", "Lancer un programme de réactivation des clients à commande unique (26 % de la base)",
     "Campagne Email de relance à J+30 avec offre ciblée Mode/Beauté (catégories à marge élevée) "
     "pour transformer un achat ponctuel en client récurrent."),
    ("5", "Prioriser la croissance sur Mode et Beauté plutôt que sur l'Électronique",
     "Marges 45-50 % et taux de retour <7,5 % : chaque dollar de budget marketing ou de stock y "
     "génère un profit net proportionnellement plus élevé que sur l'électronique."),
]
for num, title_r, desc in recos:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(f"{num}.  ")
    rn.bold = True
    rn.font.size = Pt(10.5)
    rn.font.color.rgb = RED
    rt = p.add_run(title_r)
    rt.bold = True
    rt.font.size = Pt(10.5)
    rt.font.color.rgb = GRAY
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Cm(0.5)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GRAY_LIGHT

h1("Conclusion business orientée action")
body(
    "AfriMarket dégage un chiffre d'affaires solide (2,51M$ sur 6 mois) mais une rentabilité "
    "bridée par sa dépendance à l'Électronique — catégorie à fort volume, faible marge et fort taux "
    "de retour. La priorité immédiate n'est pas de générer plus de CA, mais de protéger et augmenter "
    "le profit net : réduire les retours électronique, résoudre le problème opérationnel de Douala, "
    "et réallouer le budget marketing d'Influenceur vers Email et vers les catégories à marge élevée. "
    "La base clients, déjà fidèle à 74 %, offre un levier de croissance rentable via la réactivation "
    "ciblée des acheteurs à commande unique — une stratégie moins coûteuse et plus rentable qu'une "
    "acquisition massive de nouveaux clients."
)

doc.save(OUT)
print("Resume executif sauvegarde ->", OUT)
